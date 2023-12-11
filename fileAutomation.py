import pandas as pd
from docx import Document


def generate_student_letter(student_name, attendance_percentage):
    # Create a new Word document
    doc = Document()

    # Add a heading to the letter
    doc.add_heading("Student Attendance Update", level=1)

    # Add the body of the letter
    letter_body = (
        f"Dear Parents,\n\n"
        f"We hope this letter finds you well. We wanted to take a moment to provide you with an update on "
        f"your child, {student_name}."
        f"\n\n"
        f"We are pleased to inform you that {student_name}'s attendance has been consistently good. "
        f"Currently, the attendance percentage is {attendance_percentage}%."
        f"\n\n"
        f"Maintaining regular attendance is crucial for academic success, and we appreciate your continued support in "
        f"encouraging {student_name} to attend classes regularly."
        f"\n\n"
        f"If you have any questions or concerns regarding {student_name}'s attendance or any other matter, "
        f"please feel free to contact us."
        f"\n\n"
        f"Thank you for your cooperation and involvement in your child's education."
        f"\n\n"
        f"Sincerely,\n\n"
        f"[Your School/Institution Name]"
    )

    doc.add_paragraph(letter_body)

    # Save the document
    doc.save(f"{student_name}.docx")


def read_excel_and_print(file_path):
    df = pd.read_excel(file_path)

    # Print each row
    for index, row in df.iterrows():
        print(row.Name)
        print(row.Attendance)
        # type(row)
        student_name = row.Name
        attendance_percentage = row.Attendance

        generate_student_letter(student_name, attendance_percentage)


excel_file_path = "Attendance.xlsx"
read_excel_and_print(excel_file_path)
